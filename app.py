import streamlit as st
from groq import Groq
import os, io, re, urllib.request
import streamlit.components.v1 as components

st.set_page_config(page_title="InfographAI", page_icon="🎨", layout="wide")

api_key = os.environ.get("GROQ_API_KEY")
if not api_key:
    try: api_key = st.secrets["GROQ_API_KEY"]
    except: pass
if not api_key:
    st.error("⚠️ GROQ_API_KEY missing.")
    st.stop()

client = Groq(api_key=api_key)

# ══════════════════════════════════════════════════════════
# CSS — Light green + Thick dark green identity
# ══════════════════════════════════════════════════════════
st.markdown("""
<style>
@import url('https://fonts.googleapis.com/css2?family=Plus+Jakarta+Sans:wght@400;500;600;700;800&family=DM+Sans:wght@300;400;500;600&family=DM+Mono:wght@400;500&display=swap');

*, *::before, *::after { box-sizing: border-box; }

.stApp {
    background: #e8f5e9 !important;
    color: #1b5e20 !important;
    font-family: 'DM Sans', sans-serif !important;
    font-size: 15px !important;
}
.stApp::before {
    content: '';
    position: fixed; inset: 0; z-index: 0; pointer-events: none;
    background:
        radial-gradient(ellipse 1100px 780px at   0%   0%, rgba(165,214,167,0.65) 0%, transparent 55%),
        radial-gradient(ellipse  900px 700px at 100%   5%, rgba(200,230,201,0.52) 0%, transparent 55%),
        radial-gradient(ellipse  800px 650px at   0% 100%, rgba(165,214,167,0.45) 0%, transparent 55%),
        radial-gradient(ellipse  950px 580px at 100% 100%, rgba(129,199,132,0.38) 0%, transparent 55%);
}
.block-container {
    padding-top: 0 !important;
    padding-bottom: 3rem !important;
    max-width: 1260px !important;
    position: relative; z-index: 1;
}
#MainMenu, footer, header { visibility: hidden; }
.stDeployButton { display: none; }
::-webkit-scrollbar { width: 6px; }
::-webkit-scrollbar-track { background: #c8e6c9; }
::-webkit-scrollbar-thumb { background: #2e7d32; border-radius: 10px; }

.ig-nav {
    display: flex; align-items: center; justify-content: space-between;
    padding: 0.88rem 1.7rem; background: #1b5e20;
    border-radius: 0 0 22px 22px; box-shadow: 0 6px 26px rgba(27,94,32,0.42); margin-bottom: 0;
}
.ig-logo-wrap { display: flex; align-items: center; gap: 0.65rem; }
.ig-logo-icon {
    width: 38px; height: 38px; background: linear-gradient(135deg, #4caf50, #a5d6a7);
    border-radius: 11px; display: flex; align-items: center; justify-content: center;
    font-size: 1.05rem; box-shadow: 0 3px 10px rgba(0,0,0,0.28);
}
.ig-logo-text { font-family: 'Plus Jakarta Sans', sans-serif; font-size: 1.2rem; font-weight: 800; color: #ffffff; letter-spacing: -0.3px; }
.ig-logo-text span { color: #a5d6a7; }
.ig-nav-pills { display: flex; align-items: center; gap: 0.4rem; flex-wrap: wrap; }
.ig-pill {
    font-family: 'DM Mono', monospace; font-size: 0.6rem; letter-spacing: 0.1em; text-transform: uppercase;
    color: #c8e6c9; border: 1.5px solid rgba(165,214,167,0.35); background: rgba(255,255,255,0.12);
    padding: 0.24rem 0.72rem; border-radius: 100px; font-weight: 600;
}
.ig-status {
    display: flex; align-items: center; gap: 0.45rem; font-size: 0.8rem; font-weight: 700; color: #a5d6a7;
    background: rgba(255,255,255,0.1); border: 1.5px solid rgba(165,214,167,0.3); padding: 0.32rem 0.9rem; border-radius: 100px;
}
.ig-status-dot {
    width: 8px; height: 8px; background: #69f0ae; border-radius: 50%;
    box-shadow: 0 0 10px rgba(105,240,174,0.9); animation: sDot 2.4s ease-in-out infinite; flex-shrink: 0;
}
@keyframes sDot { 0%,100%{opacity:1;} 50%{opacity:0.22;} }

.ig-hero { text-align: center; padding: 2.8rem 0 2.2rem; }
.ig-hero-tag {
    display: inline-flex; align-items: center; gap: 0.45rem; background: #ffffff;
    border: 2.5px solid #43a047; border-radius: 100px; padding: 0.35rem 1.1rem;
    font-size: 0.77rem; font-weight: 700; color: #2e7d32; margin-bottom: 1.3rem;
    box-shadow: 0 3px 14px rgba(67,160,71,0.22);
}
.ig-hero-tag::before {
    content:''; width:8px; height:8px; background:#43a047; border-radius:50%;
    flex-shrink:0; box-shadow:0 0 8px rgba(67,160,71,0.7); animation: tagPulse 2s ease-in-out infinite;
}
@keyframes tagPulse { 0%,100%{opacity:1;transform:scale(1);} 50%{opacity:0.3;transform:scale(0.7);} }
.ig-h1 { font-family: 'Plus Jakarta Sans', sans-serif; font-size: clamp(2.2rem, 4.5vw, 3.5rem); font-weight: 800; line-height: 1.1; letter-spacing: -1px; color: #1b5e20; margin: 0; }
.ig-h1-accent {
    font-family: 'Plus Jakarta Sans', sans-serif; font-size: clamp(2.2rem, 4.5vw, 3.5rem); font-weight: 800;
    line-height: 1.1; letter-spacing: -1px;
    background: linear-gradient(90deg, #2e7d32, #388e3c, #66bb6a);
    -webkit-background-clip: text; -webkit-text-fill-color: transparent; background-clip: text; display: block; margin-bottom: 1rem;
}
.ig-hero-sub { font-size: 1rem; color: #2e7d32; font-weight: 500; max-width: 480px; margin: 0 auto; line-height: 1.75; }

.ig-card {
    background: #ffffff; border: 2px solid #c8e6c9; border-left: 6px solid #388e3c;
    border-radius: 14px; padding: 1.5rem 1.5rem 1.1rem; margin-bottom: 1rem;
    box-shadow: 0 4px 20px rgba(56,142,60,0.14), 0 1px 4px rgba(0,0,0,0.05);
}
.ig-sec-label {
    font-family: 'Plus Jakarta Sans', sans-serif; font-size: 0.68rem; font-weight: 800;
    letter-spacing: 0.13em; text-transform: uppercase; color: #2e7d32;
    margin-bottom: 0.95rem; display: flex; align-items: center; gap: 0.5rem;
}
.ig-sec-label::after { content:''; flex:1; height:2.5px; background: linear-gradient(90deg, #a5d6a7, transparent); border-radius: 2px; }

div[data-baseweb="tab-list"] {
    background: #e8f5e9 !important; border-radius: 10px !important; padding: 4px !important;
    border: 2px solid #a5d6a7 !important; gap: 3px !important; margin-bottom: 1rem !important; width: fit-content !important;
}
div[data-baseweb="tab"] {
    border-radius: 7px !important; color: #2e7d32 !important; font-weight: 700 !important;
    font-family: 'DM Sans', sans-serif !important; font-size: 0.88rem !important;
    padding: 0.42rem 1.05rem !important; transition: all 0.18s !important; min-height: unset !important;
}
div[data-baseweb="tab"]:hover { background: #c8e6c9 !important; }
div[aria-selected="true"] { background: #1b5e20 !important; color: #ffffff !important; box-shadow: 0 3px 12px rgba(27,94,32,0.42) !important; }
div[data-baseweb="tab-panel"] { background: transparent !important; padding: 0 !important; }

div[data-testid="stTextArea"] label, div[data-testid="stTextInput"] label,
div[data-testid="stSelectbox"] label, div[data-testid="stFileUploader"] label { display: none !important; }

textarea {
    background: #f1f8e9 !important; border: 2px solid #a5d6a7 !important; border-radius: 11px !important;
    color: #1b5e20 !important; font-family: 'DM Sans', sans-serif !important; font-size: 0.95rem !important;
    padding: 0.85rem 1rem !important; line-height: 1.65 !important; transition: all 0.2s !important; resize: none !important;
}
textarea:focus { border-color: #2e7d32 !important; box-shadow: 0 0 0 3px rgba(46,125,50,0.18) !important; outline: none !important; background: #ffffff !important; }
textarea::placeholder { color: #81c784 !important; }

.stTextInput input {
    background: #f1f8e9 !important; border: 2px solid #a5d6a7 !important; border-radius: 11px !important;
    color: #1b5e20 !important; font-family: 'DM Sans', sans-serif !important; font-size: 0.95rem !important;
    padding: 0.75rem 1rem !important; transition: all 0.2s !important;
}
.stTextInput input:focus { border-color: #2e7d32 !important; box-shadow: 0 0 0 3px rgba(46,125,50,0.18) !important; outline: none !important; background: #ffffff !important; }
.stTextInput input::placeholder { color: #81c784 !important; }

div[data-baseweb="select"] > div {
    background: #f1f8e9 !important; border: 2px solid #a5d6a7 !important; border-radius: 11px !important;
    color: #1b5e20 !important; font-family: 'DM Sans', sans-serif !important;
}
div[data-baseweb="select"] > div:focus-within { border-color: #2e7d32 !important; box-shadow: 0 0 0 3px rgba(46,125,50,0.18) !important; }
div[data-baseweb="select"] span { color: #1b5e20 !important; font-weight: 600 !important; }
[data-baseweb="popover"] > div { background: #ffffff !important; border: 2px solid #a5d6a7 !important; border-radius: 11px !important; box-shadow: 0 8px 24px rgba(27,94,32,0.2) !important; }
[role="option"] { color: #1b5e20 !important; font-family: 'DM Sans', sans-serif !important; padding: 0.65rem 1rem !important; }
[role="option"]:hover { background: #e8f5e9 !important; }

[data-testid="stFileUploader"] section {
    background: #f1f8e9 !important; border: 2.5px dashed #66bb6a !important;
    border-radius: 13px !important; padding: 1.4rem 1rem !important; transition: all 0.2s !important;
}
[data-testid="stFileUploader"] section:hover { border-color: #1b5e20 !important; background: #e8f5e9 !important; }
[data-testid="stFileUploader"] section p, [data-testid="stFileUploader"] section span { color: #388e3c !important; font-weight: 500 !important; }
[data-testid="stFileUploader"] button { background: #2e7d32 !important; color: #ffffff !important; border-radius: 8px !important; border: none !important; font-weight: 700 !important; }

.stButton > button {
    width: 100% !important; background: #2e7d32 !important; border: none !important;
    border-bottom: 4px solid #1b5e20 !important; border-radius: 12px !important; color: #ffffff !important;
    font-family: 'Plus Jakarta Sans', sans-serif !important; font-weight: 800 !important; font-size: 1rem !important;
    padding: 0.88rem 1.5rem !important; margin-top: 0.55rem !important; transition: all 0.18s !important;
    box-shadow: 0 4px 16px rgba(46,125,50,0.38) !important; letter-spacing: 0.01em !important;
}
.stButton > button:hover { background: #1b5e20 !important; transform: translateY(-2px) !important; box-shadow: 0 8px 24px rgba(27,94,32,0.45) !important; }
.stButton > button:active { transform: translateY(1px) !important; border-bottom-width: 2px !important; }

.stDownloadButton > button {
    background: #e8f5e9 !important; border: 2px solid #66bb6a !important; border-bottom: 3px solid #388e3c !important;
    color: #1b5e20 !important; font-family: 'DM Sans', sans-serif !important; font-weight: 700 !important;
    border-radius: 10px !important; padding: 0.5rem 1.1rem !important; font-size: 0.88rem !important; transition: all 0.18s !important;
}
.stDownloadButton > button:hover { background: #c8e6c9 !important; border-color: #2e7d32 !important; }

.sec-row {
    display: flex; align-items: center; justify-content: space-between;
    background: #f1f8e9; border: 1.5px solid #c8e6c9; border-left: 5px solid #4caf50;
    border-radius: 10px; padding: 0.62rem 1rem; margin-bottom: 0.42rem; transition: all 0.15s;
}
.sec-row:hover { background: #e8f5e9; border-left-color: #1b5e20; }
.sec-row-label { font-size: 0.88rem; font-weight: 700; color: #1b5e20; }
.badge-pass { font-size: 0.74rem; font-weight: 800; color: #1b5e20; background: #c8e6c9; border: 2px solid #66bb6a; border-radius: 6px; padding: 0.14rem 0.52rem; }
.badge-fail { font-size: 0.74rem; font-weight: 800; color: #b71c1c; background: #ffebee; border: 2px solid #ef9a9a; border-radius: 6px; padding: 0.14rem 0.52rem; }
.badge-info { font-size: 0.74rem; font-weight: 800; color: #e65100; background: #fff8e1; border: 2px solid #ffe082; border-radius: 6px; padding: 0.14rem 0.52rem; }

.score-row {
    display: flex; align-items: center; gap: 1.4rem; background: #e8f5e9;
    border: 2.5px solid #a5d6a7; border-radius: 13px; padding: 0.95rem 1.2rem; margin-top: 0.55rem;
}
.score-num { font-family:'Plus Jakarta Sans',sans-serif; font-size:2.2rem; font-weight:800; line-height:1; }
.score-lbl { font-family:'DM Mono',monospace; font-size:0.58rem; font-weight:700; letter-spacing:0.12em; text-transform:uppercase; color:#388e3c; margin-bottom:0.14rem; }
.score-bar-bg { height:10px; background:#c8e6c9; border-radius:100px; overflow:hidden; flex:1; border:1px solid #a5d6a7; }
.score-bar-fill { height:100%; border-radius:100px; }

.canvas-header {
    display: flex; align-items: center; justify-content: space-between; padding: 1rem 1.5rem;
    background: #2e7d32; border-radius: 16px 16px 0 0; box-shadow: 0 4px 18px rgba(46,125,50,0.35);
}
.canvas-title { font-family: 'Plus Jakarta Sans', sans-serif; font-size: 0.98rem; font-weight: 800; color: #ffffff; display: flex; align-items: center; gap: 0.5rem; }
.canvas-body { border: 2.5px solid #a5d6a7; border-top: none; border-radius: 0 0 16px 16px; overflow: hidden; box-shadow: 0 6px 24px rgba(56,142,60,0.14); }
.cbadge-ready { display: inline-flex; align-items: center; gap: 0.38rem; font-size: 0.76rem; font-weight: 800; color: #1b5e20; background: #a5d6a7; border: 2px solid #66bb6a; border-radius: 100px; padding: 0.26rem 0.85rem; }
.cbadge-wait { display: inline-flex; align-items: center; gap: 0.38rem; font-size: 0.76rem; font-weight: 800; color: #e65100; background: #fff8e1; border: 2px solid #ffe082; border-radius: 100px; padding: 0.26rem 0.85rem; }
.cbadge-dot { width: 7px; height: 7px; border-radius: 50%; flex-shrink: 0; }

.empty-wrap { min-height: 560px; background: #f1f8e9; display: flex; align-items: center; justify-content: center; position: relative; overflow: hidden; }
.empty-grid { position: absolute; inset: 0; pointer-events: none; background-image: linear-gradient(rgba(56,142,60,0.12) 1px, transparent 1px), linear-gradient(90deg, rgba(56,142,60,0.12) 1px, transparent 1px); background-size: 36px 36px; }
.empty-inner { position: relative; z-index: 1; text-align: center; padding: 2rem; }
.empty-emoji { font-size:4rem; display:block; margin-bottom:1.1rem; animation:float 3.2s ease-in-out infinite; }
@keyframes float { 0%,100%{transform:translateY(0)} 50%{transform:translateY(-12px)} }
.empty-h { font-family:'Plus Jakarta Sans',sans-serif; font-size:1.5rem; font-weight:800; color:#a5d6a7; margin-bottom:0.45rem; }
.empty-p { font-size:0.92rem; color:#81c784; line-height:1.7; }
.steps { display:flex; gap:0.85rem; margin-top:1.9rem; justify-content:center; flex-wrap:wrap; }
.step { background: #ffffff; border: 2px solid #a5d6a7; border-top: 5px solid #2e7d32; border-radius: 12px; padding: 0.9rem 1.2rem; text-align: center; box-shadow: 0 3px 12px rgba(56,142,60,0.14); min-width: 90px; }
.step-n { font-family:'DM Mono',monospace; font-size:0.65rem; color:#2e7d32; font-weight:800; letter-spacing:0.1em; text-transform:uppercase; margin-bottom:0.28rem; }
.step-t { font-family:'Plus Jakarta Sans',sans-serif; font-size:0.88rem; color:#1b5e20; font-weight:800; }

.chars-badge { display: inline-flex; align-items: center; gap: 0.38rem; font-family: 'DM Mono', monospace; font-size: 0.73rem; font-weight: 700; color: #1b5e20; background: #c8e6c9; border: 2px solid #66bb6a; border-radius: 8px; padding: 0.26rem 0.75rem; margin-top: 0.55rem; }

div[data-testid="stAlert"], .stSuccess > div { background: #e8f5e9 !important; border: 2px solid #a5d6a7 !important; border-left: 5px solid #2e7d32 !important; border-radius: 11px !important; color: #1b5e20 !important; }
.stError > div { background: #ffebee !important; border: 2px solid #ef9a9a !important; border-left: 5px solid #c62828 !important; border-radius: 11px !important; }
.stWarning > div { background: #fff8e1 !important; border: 2px solid #ffe082 !important; border-left: 5px solid #f57f17 !important; border-radius: 11px !important; }
.stSpinner > div { border-top-color: #1b5e20 !important; }

.export-hint { font-size: 0.85rem; color: #388e3c; line-height: 2; margin-top: 0.8rem; font-family: 'DM Sans', sans-serif; background: #f1f8e9; border: 1.5px solid #c8e6c9; border-radius: 10px; padding: 0.75rem 1rem; }
.export-hint strong { color: #1b5e20; font-weight: 800; }
</style>
""", unsafe_allow_html=True)


# ══════════════════════════════════════════════════════════
# SECURITY
# ══════════════════════════════════════════════════════════
MALWARE_SIGS = [b"cmd.exe", b"powershell", b"eval(", b"WScript", b"<script", b"javascript:"]

def check_malware(fb):
    fl = fb.lower()
    return not any(s.lower() in fl for s in MALWARE_SIGS)

def check_integrity(fb, ext):
    try:
        if ext == "pdf":
            import pypdf; pypdf.PdfReader(io.BytesIO(fb))
        elif ext == "docx":
            import docx; docx.Document(io.BytesIO(fb))
        return True
    except: return False

def fmt_size(b):
    if b < 1024: return f"{b} B"
    elif b < 1_048_576: return f"{b/1024:.1f} KB"
    else: return f"{b/1_048_576:.1f} MB"

def security_scan(fb, ext):
    m_ok = check_malware(fb)
    i_ok = check_integrity(fb, ext)
    sz   = fmt_size(len(fb))
    score = 100 - (0 if m_ok else 40) - (0 if i_ok else 30)
    clr   = "#1b5e20" if score == 100 else ("#e65100" if score >= 60 else "#b71c1c")
    bar   = ("linear-gradient(90deg,#2e7d32,#66bb6a)" if score == 100 else
             "linear-gradient(90deg,#ef6c00,#ffa726)" if score >= 60 else
             "linear-gradient(90deg,#c62828,#ef5350)")
    bp    = lambda ok: '<span class="badge-pass">✓ Pass</span>' if ok else '<span class="badge-fail">✗ Fail</span>'
    html  = f"""
<div class="ig-sec-label" style="margin-top:1.1rem;">Security Scan</div>
<div class="sec-row"><span class="sec-row-label">🛡️ Malware Signature Scan</span>{bp(m_ok)}</div>
<div class="sec-row"><span class="sec-row-label">✅ File Integrity Check</span>{bp(i_ok)}</div>
<div class="sec-row"><span class="sec-row-label">📦 File Size</span><span class="badge-info">{sz} · No limit</span></div>
<div class="score-row">
  <div><div class="score-lbl">Security</div>
       <div class="score-num" style="color:{clr};">{score}</div>
       <div style="font-size:0.6rem;color:#66bb6a;font-family:'DM Mono',monospace;">/100</div></div>
  <div style="flex:1;"><div class="score-lbl" style="margin-bottom:.3rem;">Confidence</div>
    <div class="score-bar-bg"><div class="score-bar-fill" style="width:{score}%;background:{bar};"></div></div>
    <div style="font-size:.7rem;color:{clr};font-family:'DM Mono',monospace;margin-top:.28rem;font-weight:800;">{score}% clean</div>
  </div>
</div>"""
    return html, m_ok, i_ok


# ══════════════════════════════════════════════════════════
# EXTRACTION
# ══════════════════════════════════════════════════════════
def extract_pdf(fb):
    import pypdf
    return "\n".join(p.extract_text() or "" for p in pypdf.PdfReader(io.BytesIO(fb)).pages).strip()

def extract_docx(fb):
    import docx
    return "\n".join(p.text for p in docx.Document(io.BytesIO(fb)).paragraphs).strip()

def fetch_url(url):
    req = urllib.request.Request(url, headers={'User-Agent': 'Mozilla/5.0'})
    with urllib.request.urlopen(req, timeout=15) as r:
        html = r.read().decode('utf-8', errors='ignore')
    txt = re.sub(r'<style[^>]*>.*?</style>', '', html, flags=re.DOTALL)
    txt = re.sub(r'<script[^>]*>.*?</script>', '', txt, flags=re.DOTALL)
    return re.sub(r'\s+', ' ', re.sub(r'<[^>]+>', ' ', txt)).strip()

def call_groq(system, user, max_tokens=8000, temp=0.18):
    r = client.chat.completions.create(
        model="llama-3.3-70b-versatile",
        messages=[{"role": "system", "content": system}, {"role": "user", "content": user}],
        temperature=temp, max_tokens=max_tokens)
    return r.choices[0].message.content.strip()


# ══════════════════════════════════════════════════════════
# INFOGRAPHIC CONFIG
# ══════════════════════════════════════════════════════════
TYPES = {
    "auto":       "🤖  AI Decides Best Format",
    "summary":    "📋  Key Points & Insights",
    "timeline":   "⏱  Timeline & Events",
    "stats":      "📊  Stats & Data Visualization",
    "comparison": "⚖  Comparison & Analysis",
    "process":    "🔄  Process Flow & Steps",
    "report":     "📑  Full Report Dashboard",
}

THEMES = {
    "Emerald & Green":      {"bg":"#071c12","primary":"#10b981","secondary":"#34d399","accent":"#a7f3d0","text":"#ffffff","subtext":"rgba(255,255,255,0.7)","card":"rgba(255,255,255,0.07)","border":"rgba(52,211,153,0.25)"},
    "Ocean Blue":           {"bg":"#040f1e","primary":"#3b82f6","secondary":"#60a5fa","accent":"#93c5fd","text":"#ffffff","subtext":"rgba(255,255,255,0.7)","card":"rgba(255,255,255,0.07)","border":"rgba(96,165,250,0.25)"},
    "Royal Purple":         {"bg":"#0f0720","primary":"#8b5cf6","secondary":"#a78bfa","accent":"#c4b5fd","text":"#ffffff","subtext":"rgba(255,255,255,0.7)","card":"rgba(255,255,255,0.07)","border":"rgba(167,139,250,0.25)"},
    "Sunset Fire":          {"bg":"#1a0800","primary":"#f97316","secondary":"#fb923c","accent":"#fed7aa","text":"#ffffff","subtext":"rgba(255,255,255,0.7)","card":"rgba(255,255,255,0.07)","border":"rgba(251,146,60,0.25)"},
    "Clean White & Blue":   {"bg":"#f8faff","primary":"#1e40af","secondary":"#3b82f6","accent":"#bfdbfe","text":"#0f172a","subtext":"#475569","card":"rgba(59,130,246,0.06)","border":"rgba(59,130,246,0.18)"},
    "Gold & Black":         {"bg":"#0a0800","primary":"#f59e0b","secondary":"#fbbf24","accent":"#fde68a","text":"#ffffff","subtext":"rgba(255,255,255,0.7)","card":"rgba(255,255,255,0.06)","border":"rgba(251,191,36,0.25)"},
    "Rose & Pink":          {"bg":"#1a0510","primary":"#ec4899","secondary":"#f472b6","accent":"#fbcfe8","text":"#ffffff","subtext":"rgba(255,255,255,0.7)","card":"rgba(255,255,255,0.07)","border":"rgba(244,114,182,0.25)"},
}

# ══════════════════════════════════════════════════════════
# SYSTEM PROMPT
# ══════════════════════════════════════════════════════════
SYS = """You are a world-class infographic designer who writes flawless HTML/CSS.

IRON RULES — violating any = failure:
• Output ONLY raw HTML. First character must be "<", last must be ">". Zero markdown, zero explanation.
• All CSS inside one <style> block in <head>. No external CSS files.
• Page: body margin:0, background = theme bg color. Inner wrapper: width:820px, margin:0 auto, padding:50px 44px 60px.
• Fonts: @import from Google Fonts only. Use "Sora" (700,800) for all headings, "Inter" (400,500,600) for body.
• REAL CONTENT ONLY — extract actual names, numbers, dates, facts from the document. NEVER write placeholder text.
• SVG charts: draw all charts as inline SVG. No <canvas>, no Chart.js, no external libraries.
• Animations: define @keyframes fadeUp {from{opacity:0;transform:translateY(24px)} to{opacity:1;transform:translateY(0)}}
  Apply animation:fadeUp 0.6s ease forwards with staggered animation-delay:0.1s increments per card.
• Numbers/stats: always font-size:3rem+, font-weight:800, colored in primary color.
• Cards: border-radius:16px, box-shadow:0 4px 24px rgba(0,0,0,0.18), overflow:hidden.
• Dividers between major sections: a full-width horizontal rule styled as a 1px line in accent color at 20% opacity.
• Footer: last element, centered, font-size:0.75rem, subtext color, content: "Generated by InfographAI"."""

# ══════════════════════════════════════════════════════════
# MASTER PROMPT BUILDER
# ══════════════════════════════════════════════════════════
def build_prompt(text, itype, title, theme_key):
    th = THEMES.get(theme_key, list(THEMES.values())[0])
    P  = th['primary']
    S  = th['secondary']
    A  = th['accent']
    BG = th['bg']
    TX = th['text']
    ST = th['subtext']
    CB = th['card']
    T  = f'"{title}"' if title.strip() else "generate a sharp, specific title from the document content (not generic)"

    # Shared CSS block injected into every format
    shared_css = f"""
@import url('https://fonts.googleapis.com/css2?family=Sora:wght@700;800&family=Inter:wght@400;500;600&display=swap');
*,*::before,*::after{{box-sizing:border-box;margin:0;padding:0}}
body{{background:{BG};color:{TX};font-family:'Inter',sans-serif;font-size:15px;line-height:1.6}}
.wrap{{width:820px;margin:0 auto;padding:50px 44px 60px}}
h1,h2,h3,h4{{font-family:'Sora',sans-serif;line-height:1.15}}
@keyframes fadeUp{{from{{opacity:0;transform:translateY(24px)}}to{{opacity:1;transform:translateY(0)}}}}
.fade{{opacity:0;animation:fadeUp 0.6s ease forwards}}
.pill{{display:inline-block;padding:4px 14px;border-radius:100px;font-size:0.75rem;font-weight:600;letter-spacing:0.05em;text-transform:uppercase}}
.card{{background:{CB};border-radius:16px;box-shadow:0 4px 24px rgba(0,0,0,0.18);overflow:hidden}}
.divider{{width:100%;height:1px;background:{A};opacity:0.2;margin:36px 0}}
footer{{text-align:center;font-size:0.75rem;color:{ST};margin-top:48px;padding-top:20px;border-top:1px solid rgba(255,255,255,0.08)}}
"""

    doc = f"\n\nDOCUMENT CONTENT (extract ALL real data from this):\n{text[:4800]}\n"

    # ── SUMMARY ──────────────────────────────────────────
    if itype == "summary":
        return f"""Build a KEY INSIGHTS infographic. Title: {T}.

EXACT HTML STRUCTURE TO BUILD:
<!DOCTYPE html><html><head><meta charset="UTF-8"><title>...</title><style>
{shared_css}
/* HERO */
.hero{{background:linear-gradient(135deg,{P},{S});padding:52px 44px;text-align:center;position:relative}}
.hero h1{{font-size:2.8rem;font-weight:800;color:#fff;margin-bottom:12px}}
.hero p{{font-size:1.05rem;color:rgba(255,255,255,0.82);max-width:560px;margin:0 auto 28px}}
.hero-stats{{display:flex;justify-content:center;gap:16px;flex-wrap:wrap}}
.hstat{{background:rgba(255,255,255,0.15);backdrop-filter:blur(8px);border:1px solid rgba(255,255,255,0.25);border-radius:12px;padding:10px 22px;text-align:center}}
.hstat-num{{font-family:'Sora',sans-serif;font-size:1.8rem;font-weight:800;color:#fff}}
.hstat-lbl{{font-size:0.72rem;color:rgba(255,255,255,0.75);text-transform:uppercase;letter-spacing:0.08em}}
/* INSIGHTS GRID */
.grid{{display:grid;grid-template-columns:1fr 1fr;gap:20px;margin-top:36px}}
.insight{{background:{CB};border-radius:16px;padding:24px;border-top:4px solid {P};box-shadow:0 4px 20px rgba(0,0,0,0.15)}}
.ins-num{{font-size:0.68rem;font-weight:700;color:{P};letter-spacing:0.12em;text-transform:uppercase;margin-bottom:8px}}
.ins-icon{{font-size:2rem;margin-bottom:10px;display:block}}
.ins-title{{font-family:'Sora',sans-serif;font-size:1.05rem;font-weight:700;color:{TX};margin-bottom:8px}}
.ins-body{{font-size:0.88rem;color:{ST};line-height:1.65}}
/* CALLOUT */
.callout{{background:{P};border-radius:16px;padding:32px 36px;margin-top:28px;position:relative;overflow:hidden}}
.callout-q{{font-size:5rem;font-family:'Sora',sans-serif;color:rgba(255,255,255,0.18);position:absolute;top:-10px;left:20px;line-height:1}}
.callout-text{{font-family:'Sora',sans-serif;font-size:1.25rem;font-weight:700;color:#fff;position:relative;z-index:1;line-height:1.5}}
/* BOTTOM STATS BAR */
.stats-bar{{display:grid;grid-template-columns:repeat(4,1fr);margin-top:28px;background:{CB};border-radius:16px;overflow:hidden}}
.sbar-item{{padding:24px 16px;text-align:center;border-right:1px solid rgba(255,255,255,0.08)}}
.sbar-item:last-child{{border-right:none}}
.sbar-num{{font-family:'Sora',sans-serif;font-size:2.6rem;font-weight:800;color:{S}}}
.sbar-lbl{{font-size:0.78rem;color:{ST};margin-top:4px}}
</style></head>
<body><div class="wrap">

<!-- HERO -->
<div class="hero fade" style="animation-delay:0s;border-radius:20px;overflow:hidden">
  <h1>TITLE_FROM_DOC</h1>
  <p>SUBTITLE_FROM_DOC</p>
  <div class="hero-stats">
    <div class="hstat"><div class="hstat-num">STAT1</div><div class="hstat-lbl">LABEL1</div></div>
    <div class="hstat"><div class="hstat-num">STAT2</div><div class="hstat-lbl">LABEL2</div></div>
    <div class="hstat"><div class="hstat-num">STAT3</div><div class="hstat-lbl">LABEL3</div></div>
  </div>
</div>

<!-- INSIGHTS GRID: 6–8 cards -->
<div class="grid" style="margin-top:36px">
  <!-- Repeat this block for each insight: -->
  <div class="insight fade" style="animation-delay:0.15s">
    <div class="ins-num">01</div>
    <span class="ins-icon">EMOJI</span>
    <div class="ins-title">INSIGHT HEADING</div>
    <div class="ins-body">2–3 sentence description extracted from document.</div>
  </div>
  <!-- ...more insight cards... -->
</div>

<!-- CALLOUT QUOTE -->
<div class="callout fade" style="animation-delay:0.5s">
  <div class="callout-q">"</div>
  <div class="callout-text">Most impactful sentence or finding from the document.</div>
</div>

<!-- STATS BAR -->
<div class="stats-bar fade" style="animation-delay:0.65s">
  <div class="sbar-item"><div class="sbar-num">NUM1</div><div class="sbar-lbl">LABEL</div></div>
  <div class="sbar-item"><div class="sbar-num">NUM2</div><div class="sbar-lbl">LABEL</div></div>
  <div class="sbar-item"><div class="sbar-num">NUM3</div><div class="sbar-lbl">LABEL</div></div>
  <div class="sbar-item"><div class="sbar-num">NUM4</div><div class="sbar-lbl">LABEL</div></div>
</div>

<footer>Generated by InfographAI</footer>
</div></body></html>

Now replace ALL CAPS placeholders with real content from the document below.
{doc}"""

    # ── TIMELINE ─────────────────────────────────────────
    elif itype == "timeline":
        return f"""Build a TIMELINE infographic. Title: {T}.

EXACT HTML STRUCTURE TO BUILD:
<!DOCTYPE html><html><head><meta charset="UTF-8"><title>...</title><style>
{shared_css}
.hero{{background:linear-gradient(135deg,{P},{S});padding:44px;text-align:center;border-radius:20px;margin-bottom:40px}}
.hero h1{{font-size:2.5rem;font-weight:800;color:#fff;margin-bottom:8px}}
.hero p{{color:rgba(255,255,255,0.8);font-size:1rem}}
/* TIMELINE */
.timeline{{position:relative;padding:20px 0}}
.timeline::before{{content:'';position:absolute;left:50%;top:0;bottom:0;width:3px;background:linear-gradient(to bottom,{P},{S});transform:translateX(-50%)}}
.tl-item{{display:flex;justify-content:flex-end;padding-right:calc(50% + 36px);margin-bottom:40px;position:relative}}
.tl-item:nth-child(even){{justify-content:flex-start;padding-right:0;padding-left:calc(50% + 36px)}}
.tl-dot{{position:absolute;left:50%;top:20px;width:18px;height:18px;background:{P};border:3px solid {BG};border-radius:50%;transform:translateX(-50%);box-shadow:0 0 0 4px {P}33;z-index:2}}
.tl-card{{background:{CB};border-radius:14px;padding:20px 22px;width:330px;box-shadow:0 4px 20px rgba(0,0,0,0.18);border-left:4px solid {P}}}
.tl-item:nth-child(even) .tl-card{{border-left:none;border-right:4px solid {S}}}
.tl-date{{display:inline-block;background:{P};color:#fff;font-size:0.72rem;font-weight:700;padding:3px 12px;border-radius:100px;margin-bottom:8px;letter-spacing:0.08em}}
.tl-item:nth-child(even) .tl-date{{background:{S}}}
.tl-title{{font-family:'Sora',sans-serif;font-size:1rem;font-weight:700;color:{TX};margin-bottom:6px}}
.tl-body{{font-size:0.85rem;color:{ST};line-height:1.6}}
/* SUMMARY BAR */
.span-bar{{background:{CB};border-radius:14px;padding:24px 28px;margin-top:40px}}
.span-label{{font-size:0.72rem;color:{ST};text-transform:uppercase;letter-spacing:0.1em;margin-bottom:12px}}
.span-track{{position:relative;height:8px;background:rgba(255,255,255,0.1);border-radius:100px;overflow:hidden}}
.span-fill{{height:100%;background:linear-gradient(90deg,{P},{S});border-radius:100px}}
.span-marks{{display:flex;justify-content:space-between;margin-top:8px}}
.span-mark{{font-size:0.72rem;color:{ST}}}
</style></head>
<body><div class="wrap">

<div class="hero fade"><h1>TITLE</h1><p>SUBTITLE — span of time covered</p></div>

<div class="timeline">
  <!-- Repeat for each event/milestone (minimum 6): -->
  <div class="tl-item fade" style="animation-delay:0.1s">
    <div class="tl-dot"></div>
    <div class="tl-card">
      <span class="tl-date">YEAR / DATE</span>
      <div class="tl-title">Event Title</div>
      <div class="tl-body">Description of what happened, from the document.</div>
    </div>
  </div>
  <!-- ...more timeline items... -->
</div>

<div class="span-bar fade">
  <div class="span-label">Timeline Span</div>
  <div class="span-track"><div class="span-fill" style="width:100%"></div></div>
  <div class="span-marks">
    <span class="span-mark">EARLIEST DATE</span>
    <span class="span-mark">MID POINT</span>
    <span class="span-mark">LATEST DATE</span>
  </div>
</div>

<footer>Generated by InfographAI</footer>
</div></body></html>

Replace ALL CAPS with real events, dates, and descriptions from the document below.
{doc}"""

    # ── STATS ─────────────────────────────────────────────
    elif itype == "stats":
        return f"""Build a STATS DASHBOARD infographic. Title: {T}.

EXACT HTML STRUCTURE TO BUILD:
<!DOCTYPE html><html><head><meta charset="UTF-8"><title>...</title><style>
{shared_css}
.hero{{background:{P};padding:36px 44px;border-radius:20px;display:flex;align-items:center;justify-content:space-between;margin-bottom:32px}}
.hero h1{{font-family:'Sora',sans-serif;font-size:2.2rem;font-weight:800;color:#fff}}
.hero p{{color:rgba(255,255,255,0.8);font-size:0.95rem;margin-top:6px}}
.hero-badge{{background:rgba(255,255,255,0.15);border:1px solid rgba(255,255,255,0.3);padding:8px 18px;border-radius:10px;color:#fff;font-size:0.8rem;font-weight:600;white-space:nowrap}}
/* KPI ROW */
.kpi-row{{display:grid;grid-template-columns:repeat(4,1fr);gap:16px;margin-bottom:28px}}
.kpi{{background:{CB};border-radius:14px;padding:22px 18px;border-left:4px solid {P};box-shadow:0 4px 20px rgba(0,0,0,0.15)}}
.kpi:nth-child(even){{border-left-color:{S}}}
.kpi-num{{font-family:'Sora',sans-serif;font-size:2.6rem;font-weight:800;color:{P};line-height:1}}
.kpi:nth-child(even) .kpi-num{{color:{S}}}
.kpi-lbl{{font-size:0.78rem;color:{ST};margin-top:6px;font-weight:500}}
.kpi-trend{{font-size:0.75rem;margin-top:8px;font-weight:600}}
/* CHART ROW */
.chart-row{{display:grid;grid-template-columns:1fr 1fr;gap:20px;margin-bottom:28px}}
.chart-card{{background:{CB};border-radius:16px;padding:28px;box-shadow:0 4px 20px rgba(0,0,0,0.15)}}
.chart-title{{font-family:'Sora',sans-serif;font-size:1rem;font-weight:700;color:{TX};margin-bottom:20px}}
/* BAR CHART */
.bar-row{{display:flex;align-items:center;gap:10px;margin-bottom:14px}}
.bar-label{{font-size:0.8rem;color:{ST};width:110px;flex-shrink:0;text-align:right}}
.bar-track{{flex:1;height:20px;background:rgba(255,255,255,0.08);border-radius:100px;overflow:hidden}}
.bar-fill{{height:100%;background:linear-gradient(90deg,{P},{S});border-radius:100px}}
.bar-value{{font-size:0.82rem;font-weight:700;color:{TX};width:44px;flex-shrink:0}}
/* INSIGHT CALLOUT */
.insight-box{{background:linear-gradient(135deg,{P}22,{S}22);border:1px solid {P}55;border-radius:16px;padding:28px 32px;display:flex;align-items:center;gap:24px;margin-bottom:28px}}
.insight-icon{{font-size:3rem;flex-shrink:0}}
.insight-text{{font-family:'Sora',sans-serif;font-size:1.15rem;font-weight:700;color:{TX}}}
.insight-sub{{font-size:0.88rem;color:{ST};margin-top:6px}}
</style></head>
<body><div class="wrap">

<div class="hero fade">
  <div><h1>DASHBOARD TITLE</h1><p>Subtitle / data period / source</p></div>
  <div class="hero-badge">📊 Data Report</div>
</div>

<!-- KPI ROW: 4 key metrics from the document -->
<div class="kpi-row">
  <div class="kpi fade" style="animation-delay:0.1s"><div class="kpi-num">NUM</div><div class="kpi-lbl">METRIC LABEL</div><div class="kpi-trend" style="color:#4ade80">↑ trend note</div></div>
  <div class="kpi fade" style="animation-delay:0.2s"><div class="kpi-num">NUM</div><div class="kpi-lbl">METRIC LABEL</div><div class="kpi-trend" style="color:#f87171">↓ trend note</div></div>
  <div class="kpi fade" style="animation-delay:0.3s"><div class="kpi-num">NUM</div><div class="kpi-lbl">METRIC LABEL</div><div class="kpi-trend" style="color:#4ade80">↑ trend note</div></div>
  <div class="kpi fade" style="animation-delay:0.4s"><div class="kpi-num">NUM</div><div class="kpi-lbl">METRIC LABEL</div><div class="kpi-trend" style="color:{S}">→ trend note</div></div>
</div>

<!-- CHARTS ROW -->
<div class="chart-row">
  <!-- LEFT: SVG Donut Chart -->
  <div class="chart-card fade" style="animation-delay:0.45s">
    <div class="chart-title">CHART TITLE from doc</div>
    <div style="display:flex;align-items:center;gap:24px">
      <svg width="160" height="160" viewBox="0 0 160 160">
        <circle cx="80" cy="80" r="62" fill="none" stroke="rgba(255,255,255,0.08)" stroke-width="18"/>
        <!-- Primary segment (e.g. 68%): circumference=2π×62≈390; 68%×390≈265 -->
        <circle cx="80" cy="80" r="62" fill="none" stroke="{P}" stroke-width="18" stroke-dasharray="265 390" stroke-dashoffset="98" stroke-linecap="round" transform="rotate(-90 80 80)"/>
        <!-- Secondary segment -->
        <circle cx="80" cy="80" r="62" fill="none" stroke="{S}" stroke-width="18" stroke-dasharray="78 390" stroke-dashoffset="-167" stroke-linecap="round" transform="rotate(-90 80 80)"/>
        <text x="80" y="75" text-anchor="middle" font-size="24" font-weight="800" fill="{TX}" font-family="Sora">68%</text>
        <text x="80" y="96" text-anchor="middle" font-size="11" fill="{ST}" font-family="Inter">PRIMARY</text>
      </svg>
      <div>
        <!-- Legend items -->
        <div style="display:flex;align-items:center;gap:8px;margin-bottom:10px"><div style="width:12px;height:12px;background:{P};border-radius:50%"></div><span style="font-size:0.82rem;color:{TX}">Category A: 68%</span></div>
        <div style="display:flex;align-items:center;gap:8px"><div style="width:12px;height:12px;background:{S};border-radius:50%"></div><span style="font-size:0.82rem;color:{TX}">Category B: 20%</span></div>
      </div>
    </div>
  </div>

  <!-- RIGHT: Horizontal Bar Chart -->
  <div class="chart-card fade" style="animation-delay:0.55s">
    <div class="chart-title">RANKING TITLE from doc</div>
    <!-- 5 bars: replace label/percent with real data -->
    <div class="bar-row"><span class="bar-label">Item A</span><div class="bar-track"><div class="bar-fill" style="width:85%"></div></div><span class="bar-value">85%</span></div>
    <div class="bar-row"><span class="bar-label">Item B</span><div class="bar-track"><div class="bar-fill" style="width:72%"></div></div><span class="bar-value">72%</span></div>
    <div class="bar-row"><span class="bar-label">Item C</span><div class="bar-track"><div class="bar-fill" style="width:61%"></div></div><span class="bar-value">61%</span></div>
    <div class="bar-row"><span class="bar-label">Item D</span><div class="bar-track"><div class="bar-fill" style="width:48%"></div></div><span class="bar-value">48%</span></div>
    <div class="bar-row"><span class="bar-label">Item E</span><div class="bar-track"><div class="bar-fill" style="width:35%"></div></div><span class="bar-value">35%</span></div>
  </div>
</div>

<!-- KEY INSIGHT -->
<div class="insight-box fade" style="animation-delay:0.65s">
  <div class="insight-icon">💡</div>
  <div><div class="insight-text">KEY INSIGHT: most important finding from the data.</div><div class="insight-sub">Supporting detail or context from the document.</div></div>
</div>

<footer>Generated by InfographAI</footer>
</div></body></html>

Replace ALL CAPS and example numbers with real stats from the document. Recalculate SVG stroke-dasharray values precisely.
{doc}"""

    # ── COMPARISON ────────────────────────────────────────
    elif itype == "comparison":
        return f"""Build a COMPARISON infographic. Title: {T}.

EXACT HTML STRUCTURE TO BUILD:
<!DOCTYPE html><html><head><meta charset="UTF-8"><title>...</title><style>
{shared_css}
.hero{{text-align:center;padding:44px 44px 32px;background:linear-gradient(135deg,{P}22,{S}22);border-radius:20px;margin-bottom:32px;border:1px solid {P}33}}
.hero h1{{font-size:2.4rem;font-weight:800;color:{TX};margin-bottom:8px}}
.hero p{{font-size:0.95rem;color:{ST}}}
/* VS HERO */
.vs-row{{display:grid;grid-template-columns:1fr auto 1fr;gap:20px;align-items:start;margin-bottom:32px}}
.vs-card{{background:{CB};border-radius:16px;padding:28px;box-shadow:0 4px 20px rgba(0,0,0,0.18)}}
.vs-card.left{{border-top:5px solid {P}}}
.vs-card.right{{border-top:5px solid {S}}}
.vs-name{{font-family:'Sora',sans-serif;font-size:1.4rem;font-weight:800;color:{TX};margin-bottom:4px}}
.vs-icon{{font-size:2.5rem;margin-bottom:12px;display:block}}
.vs-fact{{display:flex;align-items:flex-start;gap:8px;margin-bottom:8px;font-size:0.88rem;color:{ST}}}
.vs-fact::before{{content:'✦';color:{P};flex-shrink:0;font-size:0.75rem;margin-top:2px}}
.vs-card.right .vs-fact::before{{color:{S}}}
.vs-badge{{display:flex;align-items:center;justify-content:center;width:56px;height:56px;background:linear-gradient(135deg,{P},{S});border-radius:50%;font-family:'Sora',sans-serif;font-size:0.9rem;font-weight:800;color:#fff;box-shadow:0 4px 16px {P}66;flex-shrink:0;align-self:center}}
/* COMPARISON TABLE */
.comp-table{{width:100%;border-collapse:separate;border-spacing:0;margin-bottom:28px}}
.comp-table th{{background:{P};color:#fff;padding:14px 18px;font-size:0.82rem;font-weight:700;text-transform:uppercase;letter-spacing:0.08em}}
.comp-table th:first-child{{border-radius:10px 0 0 0}}
.comp-table th:last-child{{border-radius:0 10px 0 0;background:{S}}}
.comp-table td{{padding:13px 18px;font-size:0.88rem;border-bottom:1px solid rgba(255,255,255,0.06)}}
.comp-table tr:nth-child(even) td{{background:rgba(255,255,255,0.03)}}
.comp-table td:first-child{{font-weight:600;color:{TX}}}
.badge-yes{{background:{P}22;color:{P};border:1px solid {P}55;padding:3px 12px;border-radius:100px;font-size:0.78rem;font-weight:700}}
.badge-no{{background:#ef444422;color:#ef4444;border:1px solid #ef444455;padding:3px 12px;border-radius:100px;font-size:0.78rem;font-weight:700}}
/* VERDICT */
.verdict{{background:linear-gradient(135deg,{P},{S});border-radius:16px;padding:30px 36px;text-align:center;margin-bottom:28px}}
.verdict-label{{font-size:0.72rem;color:rgba(255,255,255,0.75);text-transform:uppercase;letter-spacing:0.15em;margin-bottom:8px}}
.verdict-text{{font-family:'Sora',sans-serif;font-size:1.3rem;font-weight:800;color:#fff;line-height:1.4}}
/* PROS CONS */
.pros-cons{{display:grid;grid-template-columns:1fr 1fr;gap:20px}}
.pc-card{{background:{CB};border-radius:14px;padding:22px;box-shadow:0 4px 16px rgba(0,0,0,0.15)}}
.pc-title{{font-family:'Sora',sans-serif;font-size:1rem;font-weight:700;margin-bottom:14px}}
.pc-item{{display:flex;gap:10px;margin-bottom:10px;font-size:0.87rem;color:{ST}}}
.check{{color:#4ade80;font-weight:700;flex-shrink:0}}
.cross{{color:#f87171;font-weight:700;flex-shrink:0}}
</style></head>
<body><div class="wrap">

<div class="hero fade"><h1>TITLE — A vs B</h1><p>Subtitle: what is being compared and why it matters</p></div>

<div class="vs-row fade" style="animation-delay:0.1s">
  <div class="vs-card left">
    <span class="vs-icon">EMOJI_A</span>
    <div class="vs-name">OPTION A NAME</div>
    <div class="vs-fact">Key fact 1 from document</div>
    <div class="vs-fact">Key fact 2 from document</div>
    <div class="vs-fact">Key fact 3 from document</div>
  </div>
  <div class="vs-badge">VS</div>
  <div class="vs-card right">
    <span class="vs-icon">EMOJI_B</span>
    <div class="vs-name">OPTION B NAME</div>
    <div class="vs-fact">Key fact 1 from document</div>
    <div class="vs-fact">Key fact 2 from document</div>
    <div class="vs-fact">Key fact 3 from document</div>
  </div>
</div>

<table class="comp-table fade" style="animation-delay:0.25s">
  <tr><th>Criteria</th><th>OPTION A</th><th>OPTION B</th></tr>
  <!-- 6-8 rows of real comparison criteria from document -->
  <tr><td>Criterion 1</td><td><span class="badge-yes">✓ Yes</span></td><td><span class="badge-no">✗ No</span></td></tr>
  <tr><td>Criterion 2</td><td>Value A</td><td>Value B</td></tr>
</table>

<div class="verdict fade" style="animation-delay:0.4s">
  <div class="verdict-label">⚖ Verdict</div>
  <div class="verdict-text">Conclusion sentence from document — which is better and why.</div>
</div>

<div class="pros-cons fade" style="animation-delay:0.5s">
  <div class="pc-card">
    <div class="pc-title" style="color:#4ade80">✅ Pros</div>
    <div class="pc-item"><span class="check">✓</span>Pro point 1 from document</div>
    <div class="pc-item"><span class="check">✓</span>Pro point 2 from document</div>
    <div class="pc-item"><span class="check">✓</span>Pro point 3 from document</div>
  </div>
  <div class="pc-card">
    <div class="pc-title" style="color:#f87171">❌ Cons</div>
    <div class="pc-item"><span class="cross">✗</span>Con point 1 from document</div>
    <div class="pc-item"><span class="cross">✗</span>Con point 2 from document</div>
    <div class="pc-item"><span class="cross">✗</span>Con point 3 from document</div>
  </div>
</div>

<footer>Generated by InfographAI</footer>
</div></body></html>

Fill ALL CAPS placeholders with real content from the document.
{doc}"""

    # ── PROCESS ───────────────────────────────────────────
    elif itype == "process":
        return f"""Build a PROCESS FLOW infographic. Title: {T}.

EXACT HTML STRUCTURE TO BUILD:
<!DOCTYPE html><html><head><meta charset="UTF-8"><title>...</title><style>
{shared_css}
.hero{{background:linear-gradient(135deg,{P},{S});border-radius:20px;padding:44px;margin-bottom:36px;display:flex;align-items:center;justify-content:space-between}}
.hero-left h1{{font-size:2.4rem;font-weight:800;color:#fff;margin-bottom:8px}}
.hero-left p{{color:rgba(255,255,255,0.82);font-size:0.95rem;max-width:460px}}
.hero-badge{{background:rgba(255,255,255,0.18);border:1px solid rgba(255,255,255,0.3);border-radius:14px;padding:16px 24px;text-align:center;flex-shrink:0}}
.hero-badge-num{{font-family:'Sora',sans-serif;font-size:2.5rem;font-weight:800;color:#fff}}
.hero-badge-lbl{{font-size:0.75rem;color:rgba(255,255,255,0.75);text-transform:uppercase;letter-spacing:0.1em}}
/* OVERVIEW FLOW SVG */
.flow-wrap{{background:{CB};border-radius:16px;padding:28px;margin-bottom:32px;box-shadow:0 4px 20px rgba(0,0,0,0.15)}}
.flow-title{{font-family:'Sora',sans-serif;font-size:0.85rem;font-weight:700;color:{ST};text-transform:uppercase;letter-spacing:0.12em;margin-bottom:20px}}
/* STEPS */
.step{{display:flex;gap:24px;margin-bottom:24px;align-items:flex-start}}
.step-num{{width:52px;height:52px;background:{P};border-radius:50%;display:flex;align-items:center;justify-content:center;font-family:'Sora',sans-serif;font-size:1.3rem;font-weight:800;color:#fff;flex-shrink:0;box-shadow:0 4px 16px {P}55}}
.step:nth-child(even) .step-num{{background:{S};box-shadow:0 4px 16px {S}55}}
.step-connector{{position:relative}}
.step-body{{background:{CB};border-radius:14px;padding:20px 22px;flex:1;border-left:4px solid {P};box-shadow:0 2px 12px rgba(0,0,0,0.12)}}
.step:nth-child(even) .step-body{{border-left-color:{S}}}
.step-title{{font-family:'Sora',sans-serif;font-size:1.05rem;font-weight:700;color:{TX};margin-bottom:6px}}
.step-desc{{font-size:0.88rem;color:{ST};line-height:1.65}}
.step-badge{{display:inline-block;margin-top:10px;background:{P}22;color:{P};border:1px solid {P}44;padding:3px 12px;border-radius:100px;font-size:0.75rem;font-weight:600}}
.step:nth-child(even) .step-badge{{background:{S}22;color:{S};border-color:{S}44}}
.arrow{{text-align:center;color:{P};font-size:1.4rem;margin:-8px 0;padding-left:26px}}
/* OUTCOMES */
.outcomes{{display:grid;grid-template-columns:repeat(3,1fr);gap:16px;margin-top:16px}}
.outcome{{background:{CB};border-radius:14px;padding:20px;text-align:center;border-top:3px solid {P};box-shadow:0 4px 16px rgba(0,0,0,0.12)}}
.outcome:nth-child(even){{border-top-color:{S}}}
.outcome-icon{{font-size:2rem;margin-bottom:10px;display:block}}
.outcome-title{{font-family:'Sora',sans-serif;font-size:0.95rem;font-weight:700;color:{TX};margin-bottom:6px}}
.outcome-desc{{font-size:0.82rem;color:{ST}}}
</style></head>
<body><div class="wrap">

<div class="hero fade">
  <div class="hero-left"><h1>PROCESS TITLE</h1><p>Brief description of the process from the document.</p></div>
  <div class="hero-badge"><div class="hero-badge-num">N</div><div class="hero-badge-lbl">Steps</div></div>
</div>

<!-- OVERVIEW FLOW (SVG horizontal flow of 4-5 phases) -->
<div class="flow-wrap fade" style="animation-delay:0.1s">
  <div class="flow-title">Process Overview</div>
  <svg width="740" height="60" viewBox="0 0 740 60" xmlns="http://www.w3.org/2000/svg">
    <!-- Phase boxes and arrows — adjust text to real phases from doc -->
    <rect x="0" y="10" width="130" height="40" rx="8" fill="{P}"/>
    <text x="65" y="35" text-anchor="middle" font-size="13" font-weight="700" fill="#fff" font-family="Sora">Phase 1</text>
    <text x="140" y="35" text-anchor="middle" font-size="18" fill="{S}" font-family="Sora">→</text>
    <rect x="155" y="10" width="130" height="40" rx="8" fill="{S}"/>
    <text x="220" y="35" text-anchor="middle" font-size="13" font-weight="700" fill="#fff" font-family="Sora">Phase 2</text>
    <text x="295" y="35" text-anchor="middle" font-size="18" fill="{P}" font-family="Sora">→</text>
    <rect x="310" y="10" width="130" height="40" rx="8" fill="{P}"/>
    <text x="375" y="35" text-anchor="middle" font-size="13" font-weight="700" fill="#fff" font-family="Sora">Phase 3</text>
    <text x="450" y="35" text-anchor="middle" font-size="18" fill="{S}" font-family="Sora">→</text>
    <rect x="465" y="10" width="130" height="40" rx="8" fill="{S}"/>
    <text x="530" y="35" text-anchor="middle" font-size="13" font-weight="700" fill="#fff" font-family="Sora">Phase 4</text>
    <text x="605" y="35" text-anchor="middle" font-size="18" fill="{P}" font-family="Sora">→</text>
    <rect x="620" y="10" width="120" height="40" rx="8" fill="{P}"/>
    <text x="680" y="35" text-anchor="middle" font-size="13" font-weight="700" fill="#fff" font-family="Sora">Phase 5</text>
  </svg>
</div>

<!-- DETAILED STEPS (minimum 5) -->
<div class="step fade" style="animation-delay:0.2s">
  <div class="step-num">1</div>
  <div class="step-body"><div class="step-title">Step Title from Document</div><div class="step-desc">Detailed description extracted from the document. 2-3 sentences.</div><span class="step-badge">⏱ Duration / Deliverable</span></div>
</div>
<div class="arrow">↓</div>
<!-- ...repeat for each step... -->

<div class="divider"></div>
<div class="outcomes fade" style="animation-delay:0.6s">
  <div class="outcome"><span class="outcome-icon">🎯</span><div class="outcome-title">Outcome 1</div><div class="outcome-desc">Result from document.</div></div>
  <div class="outcome"><span class="outcome-icon">📈</span><div class="outcome-title">Outcome 2</div><div class="outcome-desc">Result from document.</div></div>
  <div class="outcome"><span class="outcome-icon">✅</span><div class="outcome-title">Outcome 3</div><div class="outcome-desc">Result from document.</div></div>
</div>

<footer>Generated by InfographAI</footer>
</div></body></html>

Extract ALL real steps, phases, and outcomes from the document. Fill every placeholder.
{doc}"""

    # ── REPORT ────────────────────────────────────────────
    elif itype == "report":
        return f"""Build a FULL REPORT DASHBOARD infographic. Title: {T}.

EXACT HTML STRUCTURE TO BUILD:
<!DOCTYPE html><html><head><meta charset="UTF-8"><title>...</title><style>
{shared_css}
/* HERO */
.hero{{background:linear-gradient(135deg,{P},{S});padding:50px 44px;margin-bottom:32px;border-radius:20px;position:relative;overflow:hidden}}
.hero::after{{content:'';position:absolute;right:-60px;top:-60px;width:280px;height:280px;background:rgba(255,255,255,0.06);border-radius:50%}}
.hero-tag{{display:inline-block;background:rgba(255,255,255,0.2);color:#fff;padding:4px 14px;border-radius:100px;font-size:0.72rem;font-weight:600;letter-spacing:0.1em;text-transform:uppercase;margin-bottom:14px}}
.hero h1{{font-size:2.8rem;font-weight:800;color:#fff;margin-bottom:10px;position:relative;z-index:1}}
.hero-meta{{font-size:0.88rem;color:rgba(255,255,255,0.75);position:relative;z-index:1}}
/* EXEC SUMMARY */
.exec{{display:grid;grid-template-columns:1.4fr 1fr;gap:24px;margin-bottom:28px}}
.exec-text{{background:{CB};border-radius:16px;padding:28px;box-shadow:0 4px 20px rgba(0,0,0,0.15);border-left:4px solid {P}}}
.exec-text h3{{font-family:'Sora',sans-serif;font-size:1rem;font-weight:700;color:{TX};margin-bottom:12px}}
.exec-text p{{font-size:0.9rem;color:{ST};line-height:1.7}}
.exec-findings{{background:{CB};border-radius:16px;padding:24px;box-shadow:0 4px 20px rgba(0,0,0,0.15)}}
.exec-findings h3{{font-family:'Sora',sans-serif;font-size:1rem;font-weight:700;color:{TX};margin-bottom:14px}}
.finding{{display:flex;gap:12px;margin-bottom:14px;align-items:flex-start}}
.finding-num{{width:26px;height:26px;background:{P};border-radius:50%;display:flex;align-items:center;justify-content:center;font-size:0.75rem;font-weight:800;color:#fff;flex-shrink:0}}
.finding-text{{font-size:0.87rem;color:{ST};line-height:1.5}}
/* KPI ROW */
.kpi-row{{display:grid;grid-template-columns:repeat(4,1fr);gap:16px;margin-bottom:28px}}
.kpi{{background:{CB};border-radius:14px;padding:20px 16px;text-align:center;box-shadow:0 4px 16px rgba(0,0,0,0.15);border-top:3px solid {P}}}
.kpi:nth-child(2){{border-top-color:{S}}}
.kpi:nth-child(3){{border-top-color:{P}}}
.kpi:nth-child(4){{border-top-color:{S}}}
.kpi-icon{{font-size:1.6rem;margin-bottom:8px}}
.kpi-num{{font-family:'Sora',sans-serif;font-size:2.4rem;font-weight:800;color:{P}}}
.kpi:nth-child(even) .kpi-num{{color:{S}}}
.kpi-lbl{{font-size:0.75rem;color:{ST};margin-top:4px;font-weight:500}}
/* CONTENT COLS */
.content-cols{{display:grid;grid-template-columns:1fr 1fr;gap:24px;margin-bottom:28px}}
.col-card{{background:{CB};border-radius:16px;padding:26px;box-shadow:0 4px 20px rgba(0,0,0,0.15)}}
.col-card h3{{font-family:'Sora',sans-serif;font-size:1rem;font-weight:700;color:{TX};margin-bottom:16px;padding-bottom:10px;border-bottom:1px solid rgba(255,255,255,0.08)}}
.bullet{{display:flex;gap:10px;margin-bottom:10px;font-size:0.87rem;color:{ST}}}
.bullet::before{{content:'▶';color:{P};font-size:0.6rem;margin-top:3px;flex-shrink:0}}
/* DATA TABLE */
.data-table{{width:100%;border-collapse:separate;border-spacing:0;margin-bottom:28px}}
.data-table th{{background:{P};color:#fff;padding:12px 16px;font-size:0.8rem;font-weight:700;text-transform:uppercase;letter-spacing:0.08em}}
.data-table th:first-child{{border-radius:10px 0 0 0}}
.data-table th:last-child{{border-radius:0 10px 0 0}}
.data-table td{{background:{CB};padding:11px 16px;font-size:0.87rem;color:{TX};border-bottom:1px solid rgba(255,255,255,0.05)}}
.data-table tr:nth-child(even) td{{background:rgba(255,255,255,0.03)}}
/* RECOMMENDATIONS */
.recs{{display:grid;grid-template-columns:repeat(3,1fr);gap:16px}}
.rec{{background:{CB};border-radius:14px;padding:22px;box-shadow:0 4px 16px rgba(0,0,0,0.12);border-top:4px solid {P}}}
.rec:nth-child(2){{border-top-color:{S}}}
.rec-num{{font-family:'Sora',sans-serif;font-size:1.8rem;font-weight:800;color:{P};margin-bottom:8px}}
.rec:nth-child(2) .rec-num{{color:{S}}}
.rec-title{{font-family:'Sora',sans-serif;font-size:0.95rem;font-weight:700;color:{TX};margin-bottom:6px}}
.rec-desc{{font-size:0.83rem;color:{ST};line-height:1.6}}
</style></head>
<body><div class="wrap">

<div class="hero fade">
  <div class="hero-tag">📑 Report</div>
  <h1>REPORT TITLE</h1>
  <div class="hero-meta">Period: DATE_RANGE · Source: SOURCE · Type: REPORT_TYPE</div>
</div>

<div class="exec fade" style="animation-delay:0.1s">
  <div class="exec-text"><h3>Executive Summary</h3><p>3–4 sentence summary of key findings from the document.</p></div>
  <div class="exec-findings">
    <h3>Key Findings</h3>
    <div class="finding"><div class="finding-num">1</div><div class="finding-text">Finding 1 from document</div></div>
    <div class="finding"><div class="finding-num">2</div><div class="finding-text">Finding 2 from document</div></div>
    <div class="finding"><div class="finding-num">3</div><div class="finding-text">Finding 3 from document</div></div>
  </div>
</div>

<div class="kpi-row fade" style="animation-delay:0.2s">
  <div class="kpi"><div class="kpi-icon">📊</div><div class="kpi-num">NUM1</div><div class="kpi-lbl">Metric 1</div></div>
  <div class="kpi"><div class="kpi-icon">📈</div><div class="kpi-num">NUM2</div><div class="kpi-lbl">Metric 2</div></div>
  <div class="kpi"><div class="kpi-icon">🎯</div><div class="kpi-num">NUM3</div><div class="kpi-lbl">Metric 3</div></div>
  <div class="kpi"><div class="kpi-icon">💡</div><div class="kpi-num">NUM4</div><div class="kpi-lbl">Metric 4</div></div>
</div>

<div class="content-cols fade" style="animation-delay:0.3s">
  <div class="col-card">
    <h3>SECTION A TITLE</h3>
    <div class="bullet">Bullet point from document with specific data</div>
    <div class="bullet">Bullet point from document</div>
    <div class="bullet">Bullet point from document</div>
    <div class="bullet">Bullet point from document</div>
  </div>
  <div class="col-card">
    <h3>SECTION B TITLE</h3>
    <div class="bullet">Bullet point from document</div>
    <div class="bullet">Bullet point from document</div>
    <div class="bullet">Bullet point from document</div>
    <div class="bullet">Bullet point from document</div>
  </div>
</div>

<table class="data-table fade" style="animation-delay:0.4s">
  <tr><th>Category</th><th>Value</th><th>Change</th><th>Status</th></tr>
  <!-- 4-6 rows of real data from document -->
  <tr><td>Item 1</td><td>Value</td><td>+X%</td><td>Active</td></tr>
</table>

<div class="divider"></div>
<div class="recs fade" style="animation-delay:0.5s">
  <div class="rec"><div class="rec-num">01</div><div class="rec-title">Recommendation 1</div><div class="rec-desc">Description from document.</div></div>
  <div class="rec"><div class="rec-num">02</div><div class="rec-title">Recommendation 2</div><div class="rec-desc">Description from document.</div></div>
  <div class="rec"><div class="rec-num">03</div><div class="rec-title">Recommendation 3</div><div class="rec-desc">Description from document.</div></div>
</div>

<footer>Generated by InfographAI</footer>
</div></body></html>

Fill ALL placeholders with real data from the document below.
{doc}"""

    # ── AUTO ──────────────────────────────────────────────
    else:  # auto
        return f"""Analyze the document below and choose the SINGLE BEST infographic format:
- Has clear dates/events/milestones → TIMELINE
- Compares two or more things → COMPARISON
- Heavy on numbers/percentages/data → STATS DASHBOARD
- Explains steps/workflow/method → PROCESS FLOW
- General information/findings → KEY POINTS SUMMARY

Then build a COMPLETE, STUNNING HTML infographic using that format.

REQUIREMENTS:
- Title: {T}
- Theme: bg={BG}, primary={P}, secondary={S}, text={TX}, subtext={ST}, cards={CB}
- Fonts: @import Sora (800,700) + Inter (400,500,600) from Google Fonts
- Width: 820px wrapper, margin:0 auto
- Animations: @keyframes fadeUp, staggered animation-delay per section
- REAL content only — extract actual data, names, numbers from the document
- Bold big numbers (3rem+) for all stats
- SVG charts where applicable
- Cards: border-radius:16px, box-shadow:0 4px 24px rgba(0,0,0,0.18)
- Footer: "Generated by InfographAI"
{doc}"""


# ══════════════════════════════════════════════════════════
# SESSION
# ══════════════════════════════════════════════════════════
for k, v in [("doc_text", ""), ("html_out", ""), ("generated", False)]:
    if k not in st.session_state: st.session_state[k] = v


# ══════════════════════════════════════════════════════════
# NAV
# ══════════════════════════════════════════════════════════
st.markdown("""
<div class="ig-nav">
  <div class="ig-logo-wrap">
    <div class="ig-logo-icon">🎨</div>
    <div class="ig-logo-text">Infograph<span>AI</span></div>
  </div>
  <div class="ig-nav-pills">
    <span class="ig-pill">PDF</span><span class="ig-pill">DOCX</span>
    <span class="ig-pill">TEXT</span><span class="ig-pill">URL</span>
    <span class="ig-pill">7 FORMATS</span><span class="ig-pill">7 THEMES</span>
  </div>
  <div class="ig-status"><div class="ig-status-dot"></div>AI Online</div>
</div>
""", unsafe_allow_html=True)

st.markdown("""
<div class="ig-hero">
  <div class="ig-hero-tag">AI-Powered Infographic Generator</div>
  <h1 class="ig-h1">Transform Documents</h1>
  <span class="ig-h1-accent">into Visual Infographics</span>
  <p class="ig-hero-sub">Paste text, upload PDF or DOCX, or drop a URL — InfographAI turns your content into stunning, shareable visuals.</p>
</div>
""", unsafe_allow_html=True)

lcol, rcol = st.columns([1, 2.2], gap="large")

# ══════════════════════════════════════════════════════════
# LEFT
# ══════════════════════════════════════════════════════════
with lcol:

    st.markdown('<div class="ig-card">', unsafe_allow_html=True)
    st.markdown('<div class="ig-sec-label">Document Input</div>', unsafe_allow_html=True)

    t1, t2, t3, t4 = st.tabs(["📝 Text", "📄 PDF", "📃 DOCX", "🔗 URL"])

    with t1:
        txt = st.text_area("", height=180,
            placeholder="Paste your document, report, article or any text here…",
            label_visibility="collapsed", key="itxt")
        if st.button("Load Text →", key="btxt"):
            if txt.strip():
                st.session_state.doc_text = txt.strip()
                st.success(f"✅ Loaded {len(txt):,} characters")
            else:
                st.warning("Please enter some text first.")

    with t2:
        pf = st.file_uploader("", type=["pdf"], label_visibility="collapsed", key="ipdf")
        if st.button("Extract & Scan PDF →", key="bpdf"):
            if pf:
                fb = pf.read()
                sh, m_ok, i_ok = security_scan(fb, "pdf")
                st.markdown(sh, unsafe_allow_html=True)
                if m_ok and i_ok:
                    try:
                        st.session_state.doc_text = extract_pdf(fb)
                        st.success(f"✅ Extracted {len(st.session_state.doc_text):,} chars")
                    except Exception as e: st.error(f"Failed: {e}")
                else: st.error("🚫 File blocked by security scan.")
            else: st.warning("Upload a PDF first.")

    with t3:
        df = st.file_uploader("", type=["docx"], label_visibility="collapsed", key="idocx")
        if st.button("Extract & Scan DOCX →", key="bdocx"):
            if df:
                fb = df.read()
                sh, m_ok, i_ok = security_scan(fb, "docx")
                st.markdown(sh, unsafe_allow_html=True)
                if m_ok and i_ok:
                    try:
                        st.session_state.doc_text = extract_docx(fb)
                        st.success(f"✅ Extracted {len(st.session_state.doc_text):,} chars")
                    except Exception as e: st.error(f"Failed: {e}")
                else: st.error("🚫 File blocked by security scan.")
            else: st.warning("Upload a DOCX first.")

    with t4:
        ui = st.text_input("", placeholder="https://example.com/article…",
                           label_visibility="collapsed", key="iurl")
        if st.button("Fetch URL →", key="burl"):
            if ui.startswith("http"):
                with st.spinner("Fetching…"):
                    try:
                        t = fetch_url(ui)
                        if len(t) < 100: st.error("Not enough content found.")
                        else:
                            st.session_state.doc_text = t[:8000]
                            st.success(f"✅ Fetched {len(st.session_state.doc_text):,} chars")
                    except Exception as e: st.error(f"Failed: {e}")
            else: st.warning("Enter a valid URL starting with https://")

    if st.session_state.doc_text:
        st.markdown(f'<div class="chars-badge">✓ &nbsp;{len(st.session_state.doc_text):,} chars loaded</div>',
                    unsafe_allow_html=True)
    st.markdown('</div>', unsafe_allow_html=True)

    # CONFIG
    st.markdown('<div class="ig-card">', unsafe_allow_html=True)
    st.markdown('<div class="ig-sec-label">Configuration</div>', unsafe_allow_html=True)

    st.markdown('<p style="font-family:\'DM Mono\',monospace;font-size:0.68rem;font-weight:800;letter-spacing:0.12em;text-transform:uppercase;color:#2e7d32;margin:0 0 0.4rem;">Infographic Format</p>', unsafe_allow_html=True)
    itype = st.selectbox("fmt", list(TYPES.keys()), format_func=lambda k: TYPES[k],
                         label_visibility="collapsed", key="itype")

    st.markdown('<p style="font-family:\'DM Mono\',monospace;font-size:0.68rem;font-weight:800;letter-spacing:0.12em;text-transform:uppercase;color:#2e7d32;margin:0.8rem 0 0.4rem;">Color Theme</p>', unsafe_allow_html=True)
    theme = st.selectbox("thm", list(THEMES.keys()), label_visibility="collapsed", key="itheme")

    st.markdown('<p style="font-family:\'DM Mono\',monospace;font-size:0.68rem;font-weight:800;letter-spacing:0.12em;text-transform:uppercase;color:#2e7d32;margin:0.8rem 0 0.4rem;">Custom Title (optional)</p>', unsafe_allow_html=True)
    ctitle = st.text_input("", placeholder="Leave blank — AI will generate title from content…",
                           label_visibility="collapsed", key="ictitle")

    st.markdown("<div style='height:.2rem'></div>", unsafe_allow_html=True)
    gen = st.button("✨  Generate Infographic", key="bgen")
    st.markdown('</div>', unsafe_allow_html=True)

    if st.session_state.generated:
        st.markdown('<div class="ig-card">', unsafe_allow_html=True)
        st.markdown('<div class="ig-sec-label">Export</div>', unsafe_allow_html=True)

        # HTML download
        st.download_button("⬇  Download HTML",
                           data=st.session_state.html_out.encode(),
                           file_name="infographic.html", mime="text/html")

        st.markdown("""
<div class="export-hint" style="margin-top:0.65rem;">
  <strong>📸 Download as PNG:</strong><br>
  Click the floating <span style="background:#1b5e20;color:#fff;padding:1px 8px;border-radius:5px;font-size:0.78rem;font-weight:700;">⬇ Save as PNG</span>
  button inside the preview below — downloads a full high-res image automatically.
  <br><br>
  <strong>📄 Save as PDF:</strong><br>
  Open downloaded HTML in Chrome → Ctrl+P → Save as PDF
</div>""", unsafe_allow_html=True)
        st.markdown('</div>', unsafe_allow_html=True)


# ══════════════════════════════════════════════════════════
# IMAGE EXPORT INJECTOR
# Injects html2canvas + a floating PNG download button into
# the infographic HTML so it works entirely inside the iframe.
# ══════════════════════════════════════════════════════════
def inject_image_export(html: str) -> str:
    """
    Injects html2canvas into the infographic HTML.
    Key fix: ALL UI (button + overlay) is hidden BEFORE capture starts,
    then capture runs on the .ig-capture-root div (not body), 
    so overlays/spinners never appear in the final PNG.
    """
    snippet = """
<script src="https://cdnjs.cloudflare.com/ajax/libs/html2canvas/1.4.1/html2canvas.min.js"></script>
<style>
/* Wrap the entire infographic body content for clean capture */
.ig-capture-root { position: relative; }

/* Floating button — OUTSIDE capture root via fixed positioning */
#_ig_btn {
  position: fixed; bottom: 22px; right: 22px; z-index: 99999;
  background: linear-gradient(135deg, #2e7d32, #1b5e20);
  color: #fff; border: none; border-radius: 12px;
  padding: 12px 24px;
  font-family: 'Sora', 'Inter', sans-serif;
  font-size: 0.92rem; font-weight: 800; cursor: pointer;
  box-shadow: 0 4px 20px rgba(46,125,50,0.55);
  display: flex; align-items: center; gap: 8px;
  transition: transform 0.15s, box-shadow 0.15s;
  border-bottom: 3px solid #145214;
  letter-spacing: 0.01em;
}
#_ig_btn:hover { transform: translateY(-2px); box-shadow: 0 8px 28px rgba(46,125,50,0.65); }
#_ig_btn.busy { opacity: 0.5; cursor: wait; pointer-events: none; }

/* Status bar — appears AFTER capture, never during */
#_ig_status {
  display: none;
  position: fixed; bottom: 22px; left: 50%; transform: translateX(-50%);
  background: #1b5e20; color: #fff;
  padding: 10px 24px; border-radius: 100px;
  font-family: 'Inter', sans-serif; font-size: 0.88rem; font-weight: 600;
  box-shadow: 0 4px 20px rgba(0,0,0,0.35); z-index: 99999;
  white-space: nowrap;
}
</style>

<div id="_ig_status">✅ PNG downloaded!</div>
<button id="_ig_btn" onclick="_igSave()">⬇ Save as PNG</button>

<script>
function _igSave() {
  var btn    = document.getElementById('_ig_btn');
  var status = document.getElementById('_ig_status');

  // Step 1: Hide the button completely so it won't appear in screenshot
  btn.style.visibility = 'hidden';
  btn.classList.add('busy');

  // Step 2: Let the browser repaint (button gone), THEN capture
  requestAnimationFrame(function () {
    requestAnimationFrame(function () {
      // Find the best capture target: prefer .wrap / .wrapper div, fallback to body
      var target = document.querySelector('.wrap') ||
                   document.querySelector('.wrapper') ||
                   document.querySelector('main') ||
                   document.body;

      // Measure full document height for scrollable content
      var fullH = Math.max(
        document.body.scrollHeight,
        document.documentElement.scrollHeight
      );
      var fullW = Math.max(
        document.body.scrollWidth,
        document.documentElement.scrollWidth
      );

      html2canvas(target, {
        scale        : 2,
        useCORS      : true,
        allowTaint   : true,
        backgroundColor: window.getComputedStyle(document.body).backgroundColor || '#071c12',
        logging      : false,
        scrollX      : 0,
        scrollY      : 0,
        x            : 0,
        y            : 0,
        width        : target.scrollWidth,
        height       : target.scrollHeight,
        windowWidth  : fullW,
        windowHeight : fullH
      }).then(function (canvas) {
        // Trigger download
        var a = document.createElement('a');
        a.download = 'infographic.png';
        a.href = canvas.toDataURL('image/png', 1.0);
        document.body.appendChild(a);
        a.click();
        document.body.removeChild(a);

        // Restore button + show success toast
        btn.style.visibility = 'visible';
        btn.classList.remove('busy');
        btn.innerHTML = '✅ PNG Saved!';
        status.style.display = 'block';
        setTimeout(function () {
          btn.innerHTML = '⬇ Save as PNG';
          status.style.display = 'none';
        }, 3000);

      }).catch(function (err) {
        btn.style.visibility = 'visible';
        btn.classList.remove('busy');
        alert('PNG export failed: ' + err.message);
      });
    });
  });
}
</script>"""

    if '</body>' in html:
        return html.replace('</body>', snippet + '\n</body>', 1)
    return html + snippet


# ══════════════════════════════════════════════════════════
# RIGHT — Canvas
# ══════════════════════════════════════════════════════════
with rcol:
    ready = st.session_state.generated and st.session_state.html_out
    badge = (
        '<span class="cbadge-ready"><span class="cbadge-dot" style="background:#1b5e20;box-shadow:0 0 7px rgba(27,94,32,0.7);"></span>Infographic Ready</span>'
        if ready else
        '<span class="cbadge-wait"><span class="cbadge-dot" style="background:#fb8c00;box-shadow:0 0 7px rgba(251,140,0,0.7);"></span>Awaiting Input</span>'
    )
    st.markdown(f"""
<div class="canvas-header">
  <div class="canvas-title">🖼 Output Canvas</div>
  {badge}
</div>
<div class="canvas-body">""", unsafe_allow_html=True)

    if not ready:
        st.markdown("""
<div class="empty-wrap">
  <div class="empty-grid"></div>
  <div class="empty-inner">
    <span class="empty-emoji">🎨</span>
    <div class="empty-h">Your infographic appears here</div>
    <div class="empty-p">Load a document, choose format and theme,<br>then hit Generate.</div>
    <div class="steps">
      <div class="step"><div class="step-n">Step 01</div><div class="step-t">Load Doc</div></div>
      <div class="step"><div class="step-n">Step 02</div><div class="step-t">Configure</div></div>
      <div class="step"><div class="step-n">Step 03</div><div class="step-t">Generate ✨</div></div>
    </div>
  </div>
</div>""", unsafe_allow_html=True)
    else:
        components.html(inject_image_export(st.session_state.html_out), height=1400, scrolling=True)

    st.markdown('</div>', unsafe_allow_html=True)


# ══════════════════════════════════════════════════════════
# GENERATE
# ══════════════════════════════════════════════════════════
if gen:
    if not st.session_state.doc_text.strip():
        st.error("⚠️ Please load a document first — Text, PDF, DOCX or URL.")
    else:
        prompt = build_prompt(st.session_state.doc_text, itype, ctitle, theme)
        with st.spinner("✨ Building your infographic — this takes 15-25 seconds for a detailed result…"):
            try:
                raw = call_groq(SYS, prompt, max_tokens=8000, temp=0.18)
                # Clean any stray markdown fences
                raw = re.sub(r'^```html\s*', '', raw.strip())
                raw = re.sub(r'^```\s*',     '', raw.strip())
                raw = re.sub(r'```\s*$',     '', raw.strip())
                # Find the HTML start
                if not raw.strip().startswith('<!'):
                    for tag in ['<!DOCTYPE', '<html']:
                        idx = raw.find(tag)
                        if idx != -1: raw = raw[idx:]; break
                st.session_state.html_out  = raw
                st.session_state.generated = True
                st.rerun()
            except Exception as e:
                st.error(f"Generation failed: {e}")
